import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """
    Sets inner padding (margins) of a table cell in dxa (1 pt = 20 dxa).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    """
    Sets the background color of a table cell.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url, color="4F46E5", underline=True):
    """
    Adds a functional hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def format_seconds(seconds):
    if seconds is None:
        return ""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"

def create_docx(output_path, metadata, segments, target_lang=None, cuts=None, docx_trans_mode="both", timecode_mode="every"):
    """
    Creates a styled Word Document (.docx) for the transcribed clip.

    timecode_mode controls how timestamps appear in the transcript table:
      "every"   – a timestamp on every segment row (default / legacy behaviour).
      "rows"    – one row per segment, but the timestamp + speaker only appear
                  on a speaker change; same-speaker rows keep only cut labels.
      "blocks"  – consecutive same-speaker segments are merged into one block
                  with a single timestamp + speaker and the combined text.
    The "rows"/"blocks" modes only take effect when speaker labels are present.
    """
    doc = docx.Document()
    
    # Configure document layout margins (1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(metadata['clip_name'])
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate 800

    # Add subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(24)
    sub_run = subtitle_p.add_run("Transkriptionsbericht & Metadaten-Analyse")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate 500

    # Metadata Grid (built as a borderless styled table)
    has_cuts = (cuts is not None and len(cuts) > 0)
    rows_count = 5 if has_cuts else 4
    meta_table = doc.add_table(rows=rows_count, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta_table.autofit = False
    
    # Values parsing
    loc_str = "Keine GPS-Metadaten vorhanden"
    maps_link = None
    if metadata.get('gps'):
        loc_str = metadata['gps'].get('address', metadata['gps']['formatted'])
        maps_link = metadata.get('maps_link')

    if has_cuts:
        total_duration = sum(end - start for start, end in cuts)
        cuts_summary = f"{len(cuts)} Schnitte (Gesamtdauer: {format_seconds(total_duration)})"
        labels = [
            "Dateiname:",
            "Dauer:",
            "Erstellungsdatum:",
            "Schnittliste:",
            "Ort (Metadaten):"
        ]
        values = [
            metadata['clip_name'],
            metadata['duration_str'],
            metadata['creation_date'],
            cuts_summary,
            loc_str
        ]
    else:
        labels = [
            "Dateiname:",
            "Dauer:",
            "Erstellungsdatum:",
            "Ort (Metadaten):"
        ]
        values = [
            metadata['clip_name'],
            metadata['duration_str'],
            metadata['creation_date'],
            loc_str
        ]

    for idx, (label, val) in enumerate(zip(labels, values)):
        row = meta_table.rows[idx]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(1.8)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(4)
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # Slate 600
        
        # Value cell
        cell_val = row.cells[1]
        cell_val.width = Inches(4.7)
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        
        ort_idx = 4 if has_cuts else 3
        if idx == ort_idx and maps_link:
            # Add address text first
            p_val.add_run(val + " ")
            # Add Hyperlink
            add_hyperlink(p_val, "(Auf Google Maps anzeigen)", maps_link)
        else:
            p_val.add_run(val)

    # Empty paragraph spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(12)
    spacer.paragraph_format.space_after = Pt(12)
    
    # Heading Transcription
    heading_p = doc.add_paragraph()
    heading_p.paragraph_format.space_after = Pt(8)
    heading_run = heading_p.add_run("Transkript")
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate 800

    # Segments Table
    has_translation = any(seg.get('translated') for seg in segments) and (target_lang is not None)
    
    if has_translation:
        if docx_trans_mode == "translated":
            cols = 2
            headers = ["Zeitstempel", f"Übersetzung ({target_lang.upper()})"]
            col_widths = [Inches(1.2), Inches(5.3)]
        else:
            cols = 3
            headers = ["Zeitstempel", "Originaltext", f"Übersetzung ({target_lang.upper()})"]
            col_widths = [Inches(1.2), Inches(2.65), Inches(2.65)]
    else:
        cols = 2
        headers = ["Zeitstempel", "Inhalt / Transkription"]
        col_widths = [Inches(1.2), Inches(5.3)]

    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row styling
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "4F46E5") # Premium Indigo
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White

    # Precompute overlapping cut indices and boundary labels for each segment
    is_inside_cut_list = [False] * len(segments)
    boundary_labels_map = {i: [] for i in range(len(segments))}
    
    if cuts is not None and len(cuts) > 0:
        for c_idx, (cut_start, cut_end) in enumerate(cuts):
            overlapping_indices = []
            for s_idx, seg in enumerate(segments):
                seg_start = seg.get('start', 0.0)
                seg_end = seg.get('end', 0.0)
                # Overlap check
                if seg_start < cut_end and seg_end > cut_start:
                    overlapping_indices.append(s_idx)
                    is_inside_cut_list[s_idx] = True
            
            if overlapping_indices:
                in_idx = overlapping_indices[0]
                out_idx = overlapping_indices[-1]
                boundary_labels_map[in_idx].append(f"IN #{c_idx + 1}")
                boundary_labels_map[out_idx].append(f"OUT #{c_idx + 1}")

    # ------------------------------------------------------------
    # Build the list of rows to render, depending on the timecode mode.
    # The speaker-change modes ("rows"/"blocks") only apply when speakers
    # were actually detected – otherwise fall back to a timestamp per row.
    # ------------------------------------------------------------
    has_speakers = any(seg.get('speaker') for seg in segments)
    effective_mode = timecode_mode if has_speakers else "every"

    def build_columns(original_text, translated_text):
        """Maps original/translated text onto the configured content columns."""
        if has_translation:
            if docx_trans_mode == "translated":
                return [translated_text or original_text]
            return [original_text, translated_text or original_text]
        return [original_text]

    render_rows = []

    if effective_mode == "blocks":
        i = 0
        n = len(segments)
        while i < n:
            speaker = segments[i].get('speaker')
            indices = []
            while i < n and segments[i].get('speaker') == speaker:
                indices.append(i)
                i += 1

            first_seg = segments[indices[0]]
            last_seg = segments[indices[-1]]

            block_labels = []
            for idx in indices:
                block_labels.extend(boundary_labels_map[idx])
            prefix = f"[{', '.join(block_labels)}] " if block_labels else ""

            time_text = f"{prefix}[{first_seg['start_str']} - {last_seg['end_str']}]"
            if speaker:
                time_text += f"\n{speaker}"

            original_text = " ".join(
                segments[idx]['original'] for idx in indices if segments[idx].get('original')
            )
            translated_text = ""
            if has_translation:
                translated_text = " ".join(
                    (segments[idx].get('translated') or segments[idx]['original']) for idx in indices
                )

            render_rows.append({
                'time_text': time_text,
                'cols': build_columns(original_text, translated_text),
                'inside_cut': any(is_inside_cut_list[idx] for idx in indices),
                'has_boundary': len(block_labels) > 0,
            })
    else:
        _no_prev = object()
        prev_speaker = _no_prev
        for r_idx, seg in enumerate(segments):
            speaker = seg.get('speaker')
            boundary_labels = boundary_labels_map[r_idx]
            prefix = f"[{', '.join(boundary_labels)}] " if boundary_labels else ""

            if effective_mode == "rows" and speaker == prev_speaker:
                # Same speaker as the previous row: drop the timestamp/speaker,
                # but still show any cut boundary labels for this segment.
                time_text = prefix.strip()
            else:
                time_text = f"{prefix}[{seg['start_str']} - {seg['end_str']}]"
                if speaker:
                    time_text += f"\n{speaker}"
            prev_speaker = speaker

            render_rows.append({
                'time_text': time_text,
                'cols': build_columns(seg['original'], seg.get('translated')),
                'inside_cut': is_inside_cut_list[r_idx],
                'has_boundary': len(boundary_labels) > 0,
            })

    # Data Rows
    for row_idx, rr in enumerate(render_rows):
        row = table.add_row()
        row_cells = row.cells

        if rr['inside_cut']:
            bg_color = "E2F0D9"  # Soft light green for cut range
        else:
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"

        row_data = [rr['time_text']] + rr['cols']

        for i, text in enumerate(row_data):
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)

            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)

            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(9.5)
                if rr['has_boundary']:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Dark green
                else:
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate 500
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                run.font.size = Pt(10.5)

    # Save document
    doc.save(output_path)
    print(f"Report erfolgreich gespeichert unter: {output_path}")

if __name__ == '__main__':
    # Quick test generator script
    test_meta = {
        'clip_name': 'test_video.mp4',
        'duration_str': '00:01:45',
        'creation_date': '27.05.2026 12:45:00',
        'gps': {
            'lat': 52.52,
            'lon': 13.405,
            'formatted': '+52.5200, +13.4050',
            'address': 'Berlin, Deutschland'
        },
        'maps_link': 'https://www.google.com/maps/search/?api=1&query=52.52,13.405'
    }
    test_segments = [
        {'start': 0.0, 'end': 5.0, 'start_str': '00:00:00', 'end_str': '00:00:05', 'original': 'Hallo Welt, das ist eine Testaufnahme.', 'translated': 'Hello world, this is a test recording.'},
        {'start': 5.0, 'end': 12.0, 'start_str': '00:00:05', 'end_str': '00:00:12', 'original': 'Wir analysieren GPS Metadaten und transkribieren Audio offline.', 'translated': 'We analyze GPS metadata and transcribe audio offline.'}
    ]
    # Test with a cut that overlaps both segments (2.5s to 8.0s)
    create_docx('test_output.docx', test_meta, test_segments, 'en', cuts=[[2.5, 8.0]])
    create_docx('test_translated_only.docx', test_meta, test_segments, 'en', cuts=[[2.5, 8.0]], docx_trans_mode='translated')

    # Diarized test data to exercise the timecode modes
    diarized_segments = [
        {'start': 0.0, 'end': 4.0, 'start_str': '00:00:00', 'end_str': '00:00:04', 'original': 'Guten Tag, schön dass Sie da sind.', 'translated': '', 'speaker': 'Sprecher A'},
        {'start': 4.0, 'end': 8.0, 'start_str': '00:00:04', 'end_str': '00:00:08', 'original': 'Heute sprechen wir über das Projekt.', 'translated': '', 'speaker': 'Sprecher A'},
        {'start': 8.0, 'end': 12.0, 'start_str': '00:00:08', 'end_str': '00:00:12', 'original': 'Vielen Dank für die Einladung.', 'translated': '', 'speaker': 'Sprecher B'},
        {'start': 12.0, 'end': 16.0, 'start_str': '00:00:12', 'end_str': '00:00:16', 'original': 'Fangen wir gleich an.', 'translated': '', 'speaker': 'Sprecher A'},
    ]
    create_docx('test_tc_every.docx', test_meta, diarized_segments, timecode_mode='every')
    create_docx('test_tc_rows.docx', test_meta, diarized_segments, timecode_mode='rows')
    create_docx('test_tc_blocks.docx', test_meta, diarized_segments, timecode_mode='blocks')
    print("Timecode-Modus-Testdokumente erzeugt (every/rows/blocks).")

